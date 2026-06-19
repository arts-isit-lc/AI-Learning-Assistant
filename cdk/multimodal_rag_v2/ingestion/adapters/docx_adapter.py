"""DOCX adapter — extracts text, images, and tables from Word documents."""

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

from aws_lambda_powertools import Logger
from docx import Document

from ...models.data_models import ElementType, Provenance, RawElement
from ..base_adapter import BaseAdapter

if TYPE_CHECKING:
    from ...models.data_models import FileMetadata

logger = Logger(service="multimodal-rag-ingestion")


class DocxAdapter(BaseAdapter):
    """Extract text, images, and tables from DOCX files using python-docx."""

    # Heading styles that denote section boundaries
    _HEADING_STYLES = {
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
    }

    def extract(
        self, file_content: bytes, file_metadata: FileMetadata
    ) -> list[RawElement]:
        """Extract elements from a DOCX file grouped by heading sections.

        Args:
            file_content: Raw bytes of the DOCX file.
            file_metadata: Metadata about the file.

        Returns:
            List of RawElement instances with section-level provenance.
        """
        document = Document(BytesIO(file_content))
        elements: list[RawElement] = []
        current_section = "body"
        position_index = 0

        # Extract paragraphs and inline images
        for paragraph in document.paragraphs:
            try:
                para_elements, current_section = self._extract_paragraph(
                    paragraph, current_section, position_index
                )
                elements.extend(para_elements)
                position_index += len(para_elements)
            except Exception:
                logger.exception(
                    "Failed to extract paragraph",
                    extra={
                        "section": current_section,
                        "file_key": file_metadata.file_key,
                        "position_index": position_index,
                    },
                )
                continue

        # Extract tables
        for table in document.tables:
            try:
                table_element = self._extract_table(table, current_section, position_index)
                if table_element:
                    elements.append(table_element)
                    position_index += 1
            except Exception:
                logger.exception(
                    "Failed to extract table",
                    extra={
                        "section": current_section,
                        "file_key": file_metadata.file_key,
                        "position_index": position_index,
                    },
                )
                continue

        return elements

    def _extract_paragraph(
        self, paragraph, current_section: str, position_index: int
    ) -> tuple[list[RawElement], str]:
        """Extract elements from a single paragraph.

        Updates current_section if the paragraph is a heading.

        Args:
            paragraph: A python-docx Paragraph object.
            current_section: The current heading section name.
            position_index: Current position index within the document.

        Returns:
            Tuple of (extracted elements, updated section name).
        """
        elements: list[RawElement] = []

        # Check if this paragraph is a heading (update section tracking)
        style_name = paragraph.style.name if paragraph.style else "body"
        if style_name in self._HEADING_STYLES:
            heading_text = paragraph.text.strip()
            if heading_text:
                current_section = heading_text

        # Extract text content
        text = paragraph.text.strip()
        if text:
            elements.append(
                RawElement(
                    content=text,
                    element_type=ElementType.TEXT,
                    provenance=Provenance(
                        section=current_section,
                        position_index=position_index + len(elements),
                    ),
                    raw_metadata={
                        "source": "docx_paragraph",
                        "style": style_name,
                    },
                )
            )

        # Extract inline images from the paragraph
        inline_images = self._extract_inline_images(
            paragraph, current_section, position_index + len(elements)
        )
        elements.extend(inline_images)

        return elements, current_section

    def _extract_inline_images(
        self, paragraph, current_section: str, position_index: int
    ) -> list[RawElement]:
        """Extract inline images from paragraph runs.

        Args:
            paragraph: A python-docx Paragraph object.
            current_section: The current heading section name.
            position_index: Current position index.

        Returns:
            List of RawElement for each inline image found.
        """
        elements: list[RawElement] = []

        for run in paragraph.runs:
            # Check for inline shapes (images) in the run's XML
            inline_shapes = run.element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                "//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
                "//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )

            for blip in inline_shapes:
                try:
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed_id:
                        rel = paragraph.part.rels.get(embed_id)
                        if rel and hasattr(rel, "target_part"):
                            image_blob = rel.target_part.blob
                            if image_blob:
                                elements.append(
                                    RawElement(
                                        content=image_blob,
                                        element_type=ElementType.IMAGE,
                                        provenance=Provenance(
                                            section=current_section,
                                            position_index=position_index + len(elements),
                                        ),
                                        raw_metadata={
                                            "source": "docx_inline_image",
                                            "content_type": getattr(
                                                rel.target_part, "content_type", "image/unknown"
                                            ),
                                        },
                                    )
                                )
                except Exception:
                    logger.exception(
                        "Failed to extract inline image",
                        extra={"section": current_section},
                    )
                    continue

        return elements

    def _extract_table(
        self, table, current_section: str, position_index: int
    ) -> RawElement | None:
        """Extract table content as a formatted string.

        Args:
            table: A python-docx Table object.
            current_section: The current heading section name.
            position_index: Current position index.

        Returns:
            A RawElement for the table, or None if the table is empty.
        """
        rows: list[str] = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append("\t".join(cells))

        table_content = "\n".join(rows)
        if not table_content.strip():
            return None

        return RawElement(
            content=table_content,
            element_type=ElementType.TABLE,
            provenance=Provenance(
                section=current_section,
                position_index=position_index,
            ),
            raw_metadata={"source": "docx_table"},
        )
